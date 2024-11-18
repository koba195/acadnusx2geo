import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Mapping from AcadNusx to Georgian
ACADNUSX_TO_GEORGIAN_MAP = {
    'a': "ა", 'b': "ბ", 'c': "ც", 'd': "დ", 'e': "ე", 'f': "ფ", 'g': "გ", 'h': "ჰ",'i': "ი", 'j': "ჯ", 'k': "კ", 'l': "ლ", 'm': "მ",
    'n': "ნ", 'o': "ო", 'p': "პ", 'q': "ქ", 'r': "რ", 's': "ს", 't': "ტ", 'u': "უ", 'v': "ვ", 'w': "წ", 'x': "ხ",'y': "ყ", 'z': "ზ",
    'W': "ჭ", 'R': 'ღ', 'Z': "ძ", 'C': "ჩ", 'J': "ჟ", 'S': "შ", 'T': "თ"
}

def convert_text(text):
    """Convert text using the AcadNusx to Georgian mapping"""
    return ''.join([ACADNUSX_TO_GEORGIAN_MAP.get(c, c) for c in text])

def apply_font(run, font_name):
    """Apply the specified font to a run"""
    r = run._r  # access the underlying XML element
    rPr = r.find(qn('w:rPr'))  # get the <w:rPr> element (if exists)
    
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.append(rPr)

    # Set the font
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

def process_paragraph(paragraph):
    """Process each paragraph to translate and apply font"""
    for run in paragraph.runs:
        if run.font.name == 'AcadNusx':  # Check if the font is AcadNusx
            original_text = run.text
            if original_text:
                converted_text = convert_text(original_text)
                run.text = converted_text
                apply_font(run, "Sylfaen")  # Apply Sylfaen font to converted text

def process_table(table):
    """Process all cells in a table and apply translation"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                process_paragraph(paragraph)

def select_file():
    """Open a file dialog to select the input file"""
    file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
    if file_path:
        input_file.set(file_path)  # Update the input file path in the GUI
        
        # Extract the base name of the file (without the extension) for the output file
        base_name = os.path.splitext(os.path.basename(file_path))[0]
        
        # Set the default output file name as "base_name-converted"
        output_file.set(base_name + "-converted")

def convert_file():
    """Convert the file and save the output"""
    input_path = input_file.get()
    output_name = output_file.get()

    if not input_path:
        messagebox.showerror("Error", "Please select a valid file.")
        return

    # Get the directory of the input file
    input_dir = os.path.dirname(input_path)

    # Add extension if needed
    if not output_name.endswith(".docx"):
        output_name += ".docx"

    # Create the full output path in the same directory as the input file
    output_path = os.path.join(input_dir, output_name)

    try:
        # Load the document
        doc = Document(input_path)

        # Process paragraphs in the document
        for paragraph in doc.paragraphs:
            process_paragraph(paragraph)

        # Process tables in the document
        for table in doc.tables:
            process_table(table)

        # Save the modified document
        doc.save(output_path)

        messagebox.showinfo("Success", f"Document converted successfully!\nSaved as: {output_path}")

    except Exception as e:
        messagebox.showerror("Error", f"Error occurred: {e}")

# Setting up the main window
root = tk.Tk()
root.title("AcadNusx to Georgian Converter")

# Variable to hold file path
input_file = tk.StringVar()
output_file = tk.StringVar()

# Layout
frame = tk.Frame(root, padx=20, pady=20)
frame.pack(padx=10, pady=10)

label = tk.Label(frame, text="Select a Word (.docx) file to convert:")
label.pack(pady=5)

select_button = tk.Button(frame, text="Browse", command=select_file)
select_button.pack(pady=10)

file_entry = tk.Entry(frame, textvariable=input_file, width=40, state='readonly')
file_entry.pack(pady=5)

output_label = tk.Label(frame, text="Enter output file name:")
output_label.pack(pady=5)

output_entry = tk.Entry(frame, textvariable=output_file, width=40)
output_entry.pack(pady=5)

convert_button = tk.Button(frame, text="Convert", command=convert_file)
convert_button.pack(pady=20)

root.mainloop()
